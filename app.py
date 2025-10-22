import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
import google.generativeai as genai

# Load API key from Streamlit secrets
api_key = st.secrets.get("GEMINI_API_KEY")
if api_key is None:
    st.error("GEMINI_API_KEY is not set in Streamlit secrets.")
    st.stop()

# Configure the Generative AI client
genai.configure(api_key=api_key)

# Model setup
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)


def generate_ecg_report(ecg_image):
    """Send ECG image to Google Gemini and get analysis."""
    current_date = datetime.now().strftime("%Y-%m-%d")
    prompt = f"""
Analyze this ECG image and provide a detailed report with all possible information.
If any information cannot be determined, state 'Unable to determine'.
Follow this format:

**ECG REPORT**
Date: {current_date}
- Patient Info: Name, Age, Gender
- Clinical Info: Reason for ECG, History, Medications
- ECG Technical Details: Machine, Lead, Calibration, Quality
- ECG Findings: Heart Rate, Rhythm, P Waves, PR Interval, QRS Complex, QT/QTc, ST Segment, T Waves
- Axis: P Wave, QRS, T Wave
- Conduction & Morphology: Atrial, Ventricular, QRS Morphology, ST-T Changes
- Interpretation: Normal/Abnormal, Diagnosis, Comparison with previous ECG
- Conclusion & Recommendations
"""
    chat = model.start_chat(history=[])
    response = chat.send_message([prompt, ecg_image])
    return response.text


def create_doc(report_text, ecg_image):
    """Generate Word document with ECG report and image."""
    doc = Document()
    doc.add_heading("ECG ANALYSIS REPORT", 0)

    for line in report_text.split("\n"):
        if line.strip() == "":
            continue
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.strip("*"), level=1)
        elif line.startswith("-"):
            doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading("ECG Tracing:", level=1)
    if isinstance(ecg_image, BytesIO):
        doc.add_picture(ecg_image, width=Inches(6))
    else:
        image_stream = BytesIO(ecg_image.read())
        doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def main():
    st.title("ECG Analysis Insights with Google Gemini")
    st.write("Upload an ECG image and get a detailed AI-generated report.")

    ecg_file = st.file_uploader("Upload ECG Image", type=["png", "jpg", "jpeg"])

    if ecg_file:
        st.image(ecg_file, caption="Uploaded ECG", use_column_width=True)

        if st.button("Generate ECG Report"):
            with st.spinner("Analyzing ECG image..."):
                report_text = generate_ecg_report(ecg_file)
                st.session_state["report_text"] = report_text

        if "report_text" in st.session_state:
            st.header("Generated ECG Report")
            st.markdown(st.session_state["report_text"])

            doc_stream = create_doc(st.session_state["report_text"], ecg_file)
            st.download_button(
                label="Download Report as DOCX",
                data=doc_stream,
                file_name="ECG_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


if __name__ == "__main__":
    main()
